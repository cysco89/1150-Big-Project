# Creating a word document with a content of taco recipe that I will get from a specific url, on which it will give-
# various random recipes.
"""" Writing three taco recipes that I will get from the URL,"""
import requests # importing requests so I can have access to get the recipes from the URL,
from docx.shared import Inches # this will give me the ability to adjust the picture height and width,
# to start my word document, I have to import python-docx.
import docx



# first of all. Getting data from the taco server. This URL will return various random taco recipes.
# but only three would be on word document.
url = 'https://taco-1150.herokuapp.com/random/?full_taco=true'

# creating three variable names, otherwise I will end up getting duplicates of the recipes, One for each recipe.
# these variable names I will use to get the taco recipes.
taco_recipe1 = requests.get(url).json()
taco_recipe2 = requests.get(url).json()
taco_recipe3 = requests.get(url).json()

# open word document in order to work on it, and saving to a variable called 'final_project'
final_project = docx.Document()

# adding the main title of my word document, and applying the style 'title' so it appears on top of the page,
final_project.add_paragraph('Random Taco Cookbook', 'Title')

# inserting an image to the document that will appear on the first page of the document,
inserted_image = final_project.add_picture('new_image.jpg', height=Inches(5.5), width=Inches(6))

# bellow the image will go the following,

final_project.add_paragraph('Source', 'Heading 1') # will apply the heading 1 style to 'source'

# writing three list bullets, inside the double quotation mark will be the paragraph that will be written on the page,
# to have them in bullet list, style= must be entered,
final_project.add_paragraph("Image by Tai's Captures on Unsplash", style='List Bullet')
final_project.add_paragraph("Recipe from https://taco-1150.herokuapp.com/random/?full_taco=true", style='List Bullet')
final_project.add_paragraph("Code by Francisco Paredes", style='List Bullet')

# adding a page brake to start writing on the next word slide,
final_project.add_page_break()

# writing on the next page the first taco recipes,
# to have this text appearing on top of second page, page brake must be entered!

# this will be the first recipe that I would get from the URL taco recipe, which will be a random one.
recipe_title1 = final_project.add_paragraph(f"A rich {taco_recipe1['seasoning']['name']} with {taco_recipe1['condiment']['name']}, "
                            f"{taco_recipe1['mixin']['name']}, {taco_recipe1['base_layer']['name']}"
                            f" in {taco_recipe1['shell']['name']}", 'Title')

# since each recipe has five procedures, in order to print out the whole recipe, it must be written out individually!
# for easy reading the recipe, by adding a .add_heading it will print a heading style text to it,

# whichever value for the key 'name', will also be shown on the page that will be applied with the heading style.
final_project.add_heading(f'{taco_recipe1["seasoning"]["name"]}') # this 'name' key will match with the title that is above.
final_project.add_paragraph(f'{taco_recipe1["seasoning"]["recipe"]}') # this will print out the whole recipe for 'seasoning'

# for the rest is all the same code, until I get to the final procedure of the recipe, which is 'shell'.
final_project.add_heading(f'{taco_recipe1["condiment"]["name"]}')
final_project.add_paragraph(f'{taco_recipe1["condiment"]["recipe"]}')

final_project.add_heading(f'{taco_recipe1["mixin"]["name"]}')
final_project.add_paragraph(f'{taco_recipe1["mixin"]["recipe"]}')

final_project.add_heading(f'{taco_recipe1["base_layer"]["name"]}')
final_project.add_paragraph(f'{taco_recipe1["base_layer"]["recipe"]}')


final_project.add_heading(f'{taco_recipe1["shell"]["name"]}')
final_project.add_paragraph(f'{taco_recipe1["shell"]["recipe"]}')

# Once the whole recipe is in the pages printed out, I must add a page break for the next slide that will contain the next random taco recipe,
final_project.add_page_break()



# starting from here, it will be all the same code
# The only thing that that changes is the variable name from where I get the URL from, Notice the 'taco_recipe1' 'changed to taco_recipe2'-
# plus the 'recipe_title', which in this code will not do nothing since it is just an easy reading to where the next title starts.
# the rest is all the same code! As well the page break
recipe_title2 = final_project.add_paragraph(f"A rich {taco_recipe2['seasoning']['name']} with {taco_recipe2['condiment']['name']}, "
                            f"{taco_recipe2['mixin']['name']}, {taco_recipe2['base_layer']['name']}"
                            f" in {taco_recipe2['shell']['name']}", 'Title')

final_project.add_heading(f'{taco_recipe2["seasoning"]["name"]}')
final_project.add_paragraph(f'{taco_recipe2["seasoning"]["recipe"]}')

final_project.add_heading(f'{taco_recipe2["condiment"]["name"]}')
final_project.add_paragraph(f'{taco_recipe2["condiment"]["recipe"]}')

final_project.add_heading(f'{taco_recipe2["mixin"]["name"]}')
final_project.add_paragraph(f'{taco_recipe2["mixin"]["recipe"]}')

final_project.add_heading(f'{taco_recipe2["base_layer"]["name"]}')
final_project.add_paragraph(f'{taco_recipe2["base_layer"]["recipe"]}')

final_project.add_heading(f'{taco_recipe2["shell"]["name"]}')
final_project.add_paragraph(f'{taco_recipe2["shell"]["recipe"]}')
final_project.add_page_break()

# Here starts the third and final recipe that will be printed out in the word document.
recipe_title3 = final_project.add_paragraph(f"A rich {taco_recipe3['seasoning']['name']} with {taco_recipe3['condiment']['name']}, "
                            f"{taco_recipe3['mixin']['name']}, {taco_recipe3['base_layer']['name']}"
                            f" in {taco_recipe3['shell']['name']}", 'Title')


final_project.add_heading(f'{taco_recipe3["seasoning"]["name"]}')
final_project.add_paragraph(f'{taco_recipe3["seasoning"]["recipe"]}')

final_project.add_heading(f'{taco_recipe3["condiment"]["name"]}')
final_project.add_paragraph(f'{taco_recipe3["condiment"]["recipe"]}')

final_project.add_heading(f'{taco_recipe3["mixin"]["name"]}')
final_project.add_paragraph(f'{taco_recipe3["mixin"]["recipe"]}')

final_project.add_heading(f'{taco_recipe3["base_layer"]["name"]}')
final_project.add_paragraph(f'{taco_recipe3["base_layer"]["recipe"]}')

final_project.add_heading(f'{taco_recipe3["shell"]["name"]}')
final_project.add_paragraph(f'{taco_recipe3["shell"]["recipe"]}')


""" saving my Word Document """
final_project.save('taco_recipe.docx') # to save the document, .save must be entered!




